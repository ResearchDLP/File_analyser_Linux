###############################################
# ######### ~ ALEKSI File Analyser ~ ######## #
###############################################

# Documents pdf , Jpeg/jpg , png, gif, bmp, docx ,pptx ,xlsx, txt
#from text.f4 import *
import PyPDF2
import docx
import docx2txt
import sys
import os
import magic
import xlrd
import zipfile
import shutil
import platform
import warnings
#import scripts.Imageclassifier as im
#import cleaner
import importlib.util
import fitz
#from openpyxl import load_workbook
#from pptx import Presentation

def fileAnalyzer(external_file):
    result_img = ""
    result_set = list()

    # variables ######################

    test_file = external_file
    #test_file = sys.argv[1]
    #test_file = "Logo.png"

    spec_text = importlib.util.spec_from_file_location("module.name", "C:\Program Files\ALEKSI\FileAnalyzer\\text\\f4.py" )
    foo_text = importlib.util.module_from_spec(spec_text)
    spec_text.loader.exec_module(foo_text)

    spec_img = importlib.util.spec_from_file_location("module.name", "C:\Program Files\ALEKSI\FileAnalyzer\scripts\Imageclassifier.py" )
    foo_img = importlib.util.module_from_spec(spec_img)
    spec_img.loader.exec_module(foo_img)

    spec_cls = importlib.util.spec_from_file_location("module.name","C:\Program Files\ALEKSI\FileAnalyzer\cleaner.py")
    foo_cls = importlib.util.module_from_spec(spec_cls)
    spec_cls.loader.exec_module(foo_cls)


    os_name = platform.system()

    if os_name == 'Windows':
        if not os.path.exists("C:\Program Files\ALEKSI\FileAnalyzer\images"):
            os.makedirs("C:\Program Files\ALEKSI\FileAnalyzer\images")

        if not os.path.exists("C:\Program Files\ALEKSI\FileAnalyzer\Evidence"):
            os.makedirs("C:\Program Files\ALEKSI\FileAnalyzer\Evidence")

        if not os.path.exists("C:\Program Files\ALEKSI\FileAnalyzer\\unzip_dir"):
            os.makedirs("C:\Program Files\ALEKSI\FileAnalyzer\\unzip_dir")

        media_loc_excel = r"C:\Program Files\ALEKSI\FileAnalyzer\unzip_dir\xl\media\\"
        media_loc_ppt = r"C:\Program Files\ALEKSI\FileAnalyzer\unzip_dir\ppt\media\\"
        evidence_path = r"C:\Program Files\ALEKSI\FileAnalyzer\Evidence\\"

    ##################################
    if not os.path.exists("C:\Program Files\ALEKSI\FileAnalyzer\output_text"):
        os.makedirs("C:\Program Files\ALEKSI\FileAnalyzer\output_text")

    if not os.path.exists("C:\Program Files\ALEKSI\FileAnalyzer\images"):
        os.makedirs("C:\Program Files\ALEKSI\FileAnalyzer\images")


    fw = open("C:\Program Files\ALEKSI\FileAnalyzer\output_text\output.txt", "w+")

    #################################

    def copy_evid(test_file):
        shutil.copy(test_file, evidence_path)
        for filename in os.listdir(evidence_path):
            file_path = evidence_path + filename
        return file_path


    # (check) File Exists

    if os.path.isfile(test_file) == True :
        # Extension check
        def extention_check(file_name):
            extension = os.path.splitext(file_name)[1]
            extension = extension + ".NA"
            file_ext = extension.split(".")
            print(file_ext[1])
            return file_ext[1].lower()


        def text_from_pdf(file_name) :
            pdfFileObj = open(file_name, 'rb')
            pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
            pdf_num = pdfReader.numPages
            for y in range(0, pdf_num):
                pageobj = pdfReader.getPage(y)
                text = pageobj.extractText()
                fw.truncate(0)
                fw.write(str(text))
            return 0

        def text_from_doc(file_name):
            docFileObj = docx.Document(file_name)
            doc_num = len(docFileObj.paragraphs)
            for x in range(0, doc_num):
                doc_text = docFileObj.paragraphs[x].text
                fw.write(str(doc_text))
            return 0


        def text_from_ppt(test_file):
            os.system('off2txt -s ' + test_file)
            file_name = test_file.split(".")
            output = file_name[0] +'-ascii.txt'
            os.system('more '+ output)
            fw.write(str(output))
            return 0

        def text_from_xlr(test_file):
            book = xlrd.open_workbook(test_file)

            print("The number of worksheets is {0}".format(book.nsheets))
            print("Worksheet name(s): {0}".format(book.sheet_names()))
            sh = book.sheet_by_index(0)
            for rx in range(sh.nrows):
                fw.write(str(sh.row(rx)))
                #print(sh.row(rx))
            return 0

    ##################################

        def image_from_pdf(test_file):
            #path = "E:\Research\Project\websnif\images\\"
            image_path = "C:\Program Files\ALEKSI\FileAnalyzer\Images\\"
            doc = fitz.open(test_file)
            for i in range(len(doc)):
                for img in doc.getPageImageList(i):
                    xref = img[0]
                    pix = fitz.Pixmap(doc, xref)
                    if pix.n < 5:  # this is GRAY or RGB
                        pix.writePNG("p%s-%s.png" % (i, xref))
                        shutil.move("p%s-%s.png" % (i, xref) , image_path)
                    else:  # CMYK: convert to RGB first
                        pix1 = fitz.Pixmap(fitz.csRGB, pix)
                        pix1.writePNG("p%s-%s.png" % (i, xref))
                        pix1 = None
                    pix = None


        def image_from_doc(test_file):
            ABS_PATH = os.path.dirname(os.path.realpath(test_file))
            print(ABS_PATH)
            source = os.path.join(ABS_PATH)
            directory = r"C:\Program Files\ALEKSI\FileAnalyzer\Images/"
            filename = "\\" + test_file
            filename, file_extension = os.path.splitext(filename)
            filename = filename + file_extension
            directory = os.path.join(ABS_PATH, "images/")
            print(directory)
            print("Source :" + source)
            docx2txt.process("%s%s" % (source, filename), directory)


        def image_from_excel(test_file):
            filename, file_extension = os.path.splitext(test_file)
            zip_file = filename + ".zip"

            os.rename(test_file, zip_file)
            with zipfile.ZipFile(zip_file, "r") as zip_ref:
                zip_ref.extractall("C:\Program Files\ALEKSI\FileAnalyzer\\unzip_dir")
            #media_loc_excel = r"C:\Program Files\ALEKSI\FileAnalyzer\unzip_dir\xl\media\\"
            for filename in os.listdir(media_loc_excel):
                image_path = media_loc_excel + filename
                print(image_path)
                shutil.move(image_path, "C:\Program Files\ALEKSI\FileAnalyzer\\Images\\")

        def image_from_ppt(test_file):
            print("ppt image")
            filename, file_extension = os.path.splitext(test_file)
            zip_file = filename + ".zip"
            os.rename(test_file, zip_file)
            with zipfile.ZipFile(zip_file, "r") as zip_ref:
                zip_ref.extractall("unzip_dir")
            #media_loc_ppt = r"C:\Program Files\ALEKSI\FileAnalyzer\unzip_dir\ppt\media\\"
            for filename in os.listdir(media_loc_ppt):
                image_path = media_loc_ppt + filename
                print("ppt_func : " + image_path)
                shutil.move(image_path, "C:\Program Files\ALEKSI\FileAnalyzer\images\\")

        def image_only(test_file):
            shutil.copy(test_file, 'C:\Program Files\ALEKSI\FileAnalyzer\Images\\')

        def magic_num_check(file_name):
            magic_output = magic.from_file(file_name, mime=True)
            magic_ext = magic_output.split("/")
            if magic_ext[1] == "vnd.openxmlformats-officedocument.wordprocessingml.document":
                magic_ext[1] = "docx"
                return magic_ext[1].lower()
            if magic_ext[1] == "msword":
                magic_ext[1] = "doc"
                return magic_ext[1].lower()
            elif magic_ext[1] == "vnd.openxmlformats-officedocument.presentationml.presentation":
                magic_ext[1] = "pptx"
                return magic_ext[1].lower()
            elif magic_ext[1] == "vnd.openxmlformats-officedocument.spreadsheetml.sheet":
                magic_ext[1] = "xlsx"
                return magic_ext[1].lower()
            elif magic_ext[1] == "jpeg":
                magic_ext[1] = "jpg"
                return magic_ext[1].lower()
            else:
                return magic_ext[1].lower()


        def extention_comparison():
            if extention_check(test_file) == magic_num_check(test_file):
                print("☑Extentions are ok")
                if extention_check(evidence) == "pdf":
                    text_from_pdf(evidence)
                    image_from_pdf(evidence)
                elif extention_check(evidence) == "docx":
                    text_from_doc(evidence)
                    image_from_doc(evidence)
                elif extention_check(evidence) == "pptx":
                    text_from_ppt(evidence)
                    image_from_ppt(evidence)
                elif extention_check(evidence) == "xlsx":
                    text_from_xlr(evidence)
                    image_from_excel(evidence)
                elif extention_check(evidence) == "jpg":
                    image_only(evidence)

                return 0
            else:
                return 1


        def text_analyser():
            fr = open("C:\Program Files\ALEKSI\FileAnalyzer\output_text\output.txt", "r")
            if fr.mode == 'r':
                contents = fr.read()
            cl = foo_text.Classifier()
            return cl.classifer(contents)


        def image_analyser():
            print("Image Analyzer Code accessed")
            image_path = 'C:\Program Files\ALEKSI\FileAnalyzer\Images\\'
            #graph_path = ' "scripts\\tf_files\\retrained_graph.pb"'
            #label_path =  ' "scripts\\tf_files\\retrained_labels.txt"'
            #print(graph_path)
            #print (label_path)

            numfile = len([f for f in os.listdir(image_path) if f[0] != '.'])
            count = 0
            output  = [[0 for j in range(5)] for i in range(numfile)]

            category = ""
            score = 0

            for filename in os.listdir(image_path):
                file_path = image_path + filename
                #os.chdir("/home/avishka/Desktop/extract/websnif/scripts")
                #cmd = 'python "C:\Program Files\ALEKSI\FileAnalyzer\scripts\classifier.py" ' + file_path + graph_path + label_path
                #print(cmd)
                #out_img = os.system(cmd)
                output[count][0] = filename
                output[count][1], output[count][2], output[count][3] = foo_img.image_classifier(file_path)

                if output[count][1] != "others":
                    if score < output[count][3]:

                        score = output[count][3]
                        category = output[count][1]

                print (output[count][0], output[count][1], output[count][2], output[count][3])

                count += 1

                result = str(score) + str(category)

            print("Score : " + str(score))
            print("Category : " + str(category))


            output_text = text_analyser()
            result_set.append([score, category, output_text])
            return result_set


        evidence = copy_evid(test_file)

        t1 = extention_check(evidence)
        print(t1)

        t2 = magic_num_check(evidence)
        print(t2)

        extention_check_return = extention_comparison()

        result_set.append([extention_check_return])
        result_img = image_analyser()


        #print(text_analyser())

        foo_cls.folderCleaner()

    else:
        print("Input Error : File doesnt exists")

    fw.close()

    return result_img

