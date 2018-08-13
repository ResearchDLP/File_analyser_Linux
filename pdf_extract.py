###############################################
# ######### ~ ALEKSI File Analyser ~ ######## #
###############################################

# Documents pdf , Jpeg/jpg , png, gif, bmp, docx ,pptx ,xlsx, txt
from text.f4 import *
import PyPDF2
import docx
import docx2txt
import sys
import os
import magic
import xlrd
import zipfile
import shutil
import warnings
#import fitz
#from openpyxl import load_workbook
#from pptx import Presentation

# variables ######################

#test_file = sys.argv[1]
test_file = "ppt_test.pptx"

##################################

fw = open("output_text/output.txt", "w+")

#################################

# (check) File Exists

if os.path.isfile(test_file) == True :
    # Extension check
    def extention_check(file_name):
        extension = os.path.splitext(file_name)[1]
        file_ext = extension.split(".")
        return file_ext[1].lower()


    def text_from_pdf(file_name) :
        pdfFileObj = open(file_name, 'rb')
        pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
        pdf_num = pdfReader.numPages
        for y in range(0, pdf_num):
            pageobj = pdfReader.getPage(y)
            text = pageobj.extractText()
            fw.truncate(0)
            fw.write(text)
        return 0

    def text_from_doc(file_name):
        docFileObj = docx.Document(file_name)
        doc_num = len(docFileObj.paragraphs)
        for x in range(0, doc_num):
            doc_text = docFileObj.paragraphs[x].text
            print(doc_text)
        return 0


    def text_from_ppt(test_file):
        os.system('off2txt -s ' + test_file)
        file_name = test_file.split(".")
        output = file_name[0] +'-ascii.txt'
        os.system('more '+ output)


    def text_from_xlr(test_file):
        book = xlrd.open_workbook("Book1.xlsx")

        print("The number of worksheets is {0}".format(book.nsheets))
        print("Worksheet name(s): {0}".format(book.sheet_names()))
        sh = book.sheet_by_index(0)
        for rx in range(sh.nrows):
            print(sh.row(rx))

##################################

    def image_from_pdf(test_file):
        path = "E:\Research\Project\websnif\images\\"
        doc = fitz.open(test_file)
        for i in range(len(doc)):
            for img in doc.getPageImageList(i):
                xref = img[0]
                pix = fitz.Pixmap(doc, xref)
                if pix.n < 5:  # this is GRAY or RGB
                    pix.writePNG(path + test_file + "_"+"p%s-%s.png" % (i, xref))
                else:  # CMYK: convert to RGB first
                    pix1 = fitz.Pixmap(fitz.csRGB, pix)
                    pix1.writePNG(path + test_file + "_"+"p%s-%s.png" % (i, xref))
                    pix1 = None
                pix = None


    def image_from_doc(test_file):
        ABS_PATH = os.path.dirname(os.path.realpath(test_file))
        print(ABS_PATH)
        source = os.path.join(ABS_PATH)
        directory = r"E:\Research\Project\websnif\images/"
        filename = "\\" + test_file
        filename, file_extension = os.path.splitext(filename)
        filename = filename + file_extension
        directory = os.path.join(ABS_PATH, "images/")
        print(directory)
        print("Source :" + source)
        docx2txt.process("%s%s" % (source, filename), directory)


    def image_from_excel(test_file):
        filename, file_extension = os.path.splitext(test_file)
        zip_file = filename + ".zip"
        os.rename(test_file, zip_file)
        with zipfile.ZipFile(zip_file, "r") as zip_ref:
            zip_ref.extractall("unzip_dir")
        media_loc = r"/home/avishka/Desktop/extract/websnif/unzip_dir/xl/media/"
        for filename in os.listdir(media_loc):
            image_path = media_loc + filename
            print(image_path)
            shutil.move(image_path, "/home/avishka/Desktop/extract/websnif/images")

    def image_from_ppt(test_file):
        print("ppt image")
        filename, file_extension = os.path.splitext(test_file)
        zip_file = filename + ".zip"
        os.rename(test_file, zip_file)
        with zipfile.ZipFile(zip_file, "r") as zip_ref:
            zip_ref.extractall("unzip_dir")
        media_loc = r"/home/avishka/Desktop/extract/websnif/unzip_dir/ppt/media/"
        for filename in os.listdir(media_loc):
            image_path = media_loc + filename
            print("ppt_func : " + image_path)
            shutil.move(image_path, "/home/avishka/Desktop/extract/websnif/images")

    def magic_num_check(file_name):
        magic_output = magic.from_file(file_name, mime=True)
        magic_ext = magic_output.split("/")
        if magic_ext[1] == "vnd.openxmlformats-officedocument.wordprocessingml.document":
            magic_ext[1] = "docx"
            return magic_ext[1].lower()
        if magic_ext[1] == "msword":
            magic_ext[1] = "doc"
            return magic_ext[1].lower()
        elif magic_ext[1] == "vnd.openxmlformats-officedocument.presentationml.presentation":
            magic_ext[1] = "pptx"
            return magic_ext[1].lower()
        elif magic_ext[1] == "vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            magic_ext[1] = "xlsx"
            return magic_ext[1].lower()
        elif magic_ext[1] == "jpeg":
            magic_ext[1] = "jpg"
            return magic_ext[1].lower()
        else:
            return magic_ext[1].lower()


    def extention_comparison():
        if extention_check(test_file) == magic_num_check(test_file):
            print("☑Extentions are ok")
            if extention_check(test_file) == "pdf":
                text_from_pdf(test_file)
                image_from_pdf(test_file)
            elif extention_check(test_file) == "docx":
                text_from_doc(test_file)
                image_from_doc(test_file)
            elif extention_check(test_file) == "pptx":
                text_from_ppt(test_file)
                image_from_ppt(test_file)
            elif extention_check(test_file) == "xlsx":
                text_from_xlr(test_file)
                image_from_excel(test_file)
            return 0
        else:
            print(" ALERT : File type mismatch ")
            return 1


    def text_analyser():

        text_val = "Nigerian prince"
        print("Text analyser")
        cl = Classifier()
        return  cl.classifer(text_val)




    def image_analyser():

        image_path = "/home/avishka/Desktop/extract/websnif/images/"
        graph_path = ' /home/avishka/Desktop/extract/scripts/tf_files/retrained_graph.pb'
        label_path =  ' /home/avishka/Desktop/extract/scripts/tf_files/retrained_labels.txt'
        for filename in os.listdir(image_path):
            file_path = image_path + filename
            print(file_path)
            #os.chdir("/home/avishka/Desktop/extract/websnif/scripts")
            cmd = 'python /home/avishka/Desktop/extract/scripts/classifier.py ' + file_path + graph_path + label_path + " 2> /dev/null"
            #print(cmd)
            out_img = os.system(cmd)

        return 0


    image_analyser()
    print(text_analyser())

    t1 = extention_check(test_file)
    print(t1)

    t2 = magic_num_check(test_file)
    print(t2)

    extention_comparison()

else:
    print("Input Error : File doesnt exists")

fw.close()
